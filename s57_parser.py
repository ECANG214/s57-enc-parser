#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
S-57 ENC File Parser - Group 3 (第3小组)
=========================================
Parses ISO 8211 format S-57 ENC binary files (like CN335002.000).

Usage:
    python 第3组.py                          → prompts for file path
    python 第3组.py  path/to/CN335002.000    → direct parse

Compliant with: ISO 8211:1994, IHO S-57 Edition 3.1
Requires: Python 3.7+, python-docx (pip install python-docx)
"""

import os, sys, struct

# ============================================================
# CONFIGURATION
# ============================================================
TARGET_AGENCY = 70       # CN = 70
TARGET_FIDN   = 13753099  # 0013753099
TARGET_FIDS   = 22501     # 22501
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

def get_enc_file_path():
    """Get ENC file path — command-line argument or manual input."""
    # 1. Command-line argument
    if len(sys.argv) > 1:
        path = sys.argv[1].strip('"').strip("'")
        if os.path.isfile(path):
            print(f"Using: {path}")
            return path
        print(f"[WARNING] File not found: {path}")

    # 2. Manual input
    print()
    print("=" * 60)
    print("  S-57 ENC File Parser - Group 3 (第3小组)")
    print("=" * 60)
    print()
    print("Drag-and-drop supported: python 第3组.py path/to/file.000")
    print()
    while True:
        try:
            user_input = input("Enter path to .000 file: ").strip().strip('"').strip("'")
            if not user_input:
                print("No input. Exiting.")
                sys.exit(1)
            if os.path.isfile(user_input):
                return user_input
            print(f"[WARNING] File not found: {user_input}")
            print("Please check the path and try again, or press Ctrl+C to exit.")
        except (EOFError, KeyboardInterrupt):
            print("\nExiting.")
            sys.exit(1)

ENC_FILE = None  # Set at runtime

# ============================================================
# S-57 KNOWLEDGE BASE
# ============================================================
RCNM_NAMES = {
    10:"DS General Information", 20:"DS Geographic Reference",
    100:"Feature Record", 110:"Vector - Isolated Node",
    120:"Vector - Connected Node", 130:"Vector - Edge",
    210:"Feature-Spatial Association",
}

OBJL_NAMES = {
    1:"Administration Area",2:"Anchor Berth",3:"Anchor Zone",4:"Anchorage area",
    5:"Beacon, cardinal",6:"Beacon, isolated danger",7:"Beacon, lateral",
    8:"Beacon, safe water",9:"Beacon, special purpose",10:"Berth",11:"Bridge",
    12:"Building, religious",13:"Building, single",14:"Built-up area",
    15:"Buoy, cardinal",16:"Buoy, installation",17:"Buoy, isolated danger",
    18:"Buoy, lateral",19:"Buoy, safe water",20:"Buoy, special purpose",
    21:"Cable area",22:"Cable, overhead",23:"Cable, submarine",24:"Canal",
    25:"Causeway",26:"Caution area",27:"Checkpoint",28:"Coast guard station",
    29:"Coastline",30:"Contiguous zone",31:"Continental shelf area",
    32:"Control Zone",33:"Conveyor",34:"Crane",35:"Current - non-gravitational",
    36:"Custom zone",37:"Dam",38:"Daymark",39:"Deep water route centerline",
    40:"Deep water route part",41:"Depth area",42:"Depth contour",
    43:"Distance mark",44:"Dock area",45:"Dredged area",46:"Dumping ground",
    47:"Dyke",48:"Exclusion zone",49:"Fairway",50:"Fence/wall",
    51:"Ferry route",52:"Fishing facility",53:"Fishing ground",54:"Fog signal",
    55:"Fortified structure",56:"Gat",57:"Gridiron",
    58:"Harbour area (administrative)",59:"Harbour facility",60:"Hulk",
    61:"Incineration area",62:"Information area",63:"Inshore traffic zone",
    64:"Lake",65:"Land area",66:"Land elevation",67:"Land region",
    68:"Landmark",69:"Light",70:"Local magnetic anomaly",71:"Lock basin",
    72:"Log pond",73:"Magnetic variation",74:"Marine farm/culture",
    75:"Military practice area",76:"Mooring/Warping facility",
    77:"Navigation line",78:"Obstruction",79:"Offshore platform",
    80:"Offshore production area",81:"Oil barrier",82:"Pile",
    83:"Pilot boarding place",84:"Pipeline area",85:"Pipeline, overhead",
    86:"Pipeline, submarine/on land",87:"Pontoon",88:"Precautionary area",
    89:"Production installation",90:"Pylon/bridge support",91:"Radar line",
    92:"Radar range",93:"Radar station",94:"Radio calling-in point",
    95:"Radio station",96:"Railway",97:"Rapids",98:"Recommended route centerline",
    99:"Recommended track",100:"Recommended traffic lane part",
    101:"Rescue station",102:"Restricted area",103:"River",104:"Road",
    105:"Runway",106:"Sand waves",107:"Sea area/named water area",
    108:"Sea-plane landing area",109:"Seabed area",110:"Shoreline construction",
    111:"Signal station, traffic",112:"Signal station, warning",113:"Silo/tank",
    114:"Sloping ground",115:"Sludge",116:"Small craft facility",
    117:"Sound signal",118:"Spoil ground",119:"Spring",120:"Submarine volcano",
    121:"Swept Area",122:"Tidal stream - harmonic prediction",
    123:"Tide - harmonic prediction",124:"Tide - time series",125:"Tideway",
    126:"Topmark",127:"Traffic separation line",
    128:"Traffic separation scheme boundary",
    129:"Traffic separation scheme crossing",
    130:"Traffic separation scheme lane part",
    131:"Traffic separation scheme roundabout",132:"Traffic separation zone",
    133:"Tunnel",134:"Two-way route part",135:"Underwater/awash rock",
    136:"Unsurveyed area",137:"Vegetation",138:"Water turbulence",
    139:"Waterfall",140:"Weed/Kelp",141:"Wreck",142:"Slope topline",
    143:"Nautical publication information",144:"Tidal stream panel data",
    145:"Tide panel data",146:"M_NSYS",147:"M_QUAL",148:"M_ACCY",
    149:"Cartographic line",150:"Cartographic symbol",151:"Compass",
    152:"Text",153:"Cartographic area",154:"Geographic line",
    155:"Data coverage",156:"Nautical publication information",
    157:"Quality of data",158:"Accuracy of data",159:"M_COVR (Coverage)",
    160:"M_NSYS",
}

ATTR_NAMES = {
    40:"OBJNAM - Object name",41:"NOBJNM - Object name in national language",
    42:"INFORM - Information",43:"NINFOM - Information in national language",
    46:"SCAMIN - Scale minimum",47:"SCAMAX - Scale maximum",
    48:"HORDAT - Horizontal datum",49:"VERDAT - Vertical datum",
    50:"SNDDAT - Sounding datum",51:"SNDACC - Sounding accuracy",
    52:"POSACC - Positional accuracy",53:"QUAPOS - Quality of position",
    54:"TECSOU - Technique of sounding measurement",
    56:"CATCOA - Category of coastline",58:"ELEVAT - Elevation",
    59:"HEIGHT - Height",60:"VERLEN - Vertical length",
    61:"HORLEN - Horizontal length",62:"HORWID - Horizontal width",
    63:"VERACC - Vertical accuracy",64:"HORACC - Horizontal accuracy",
    65:"WATLEV - Water level effect",70:"COLOUR - Colour",
    71:"CATWRK - Category of wreck",72:"CONDTN - Condition",
    73:"CONRAD - Conspicuous, radar",74:"CONVIS - Conspicuous, visually",
    75:"CONST - Construction material",76:"DATEND - Date end",
    77:"DATSTA - Date start",78:"EXCLIT - Exhibition condition of object",
    79:"FUNCTN - Function",80:"HUNITS - Height units",
    81:"MARSYS - Navigational mark system",82:"NATCON - Nature of construction",
    83:"NATSUR - Nature of surface",84:"PEREND - Periodic date end",
    85:"PERSTA - Periodic date start",86:"PRODCT - Product",
    87:"RADRFL - Radar conspicuous",88:"RADWAL - Radar wave length",
    89:"COMCHA - Communication channel",90:"RADIUS - Radius",
    91:"RESTRN - Restriction",92:"DATEND - Season end date",
    93:"DATSTA - Season start date",94:"SIGFRQ - Signal frequency",
    95:"SIGGEN - Signal generation",96:"SIGGRP - Signal group",
    97:"SIGPER - Signal period",98:"SIGSEQ - Signal sequence",
    99:"SORDAT - Source date",100:"SORIND - Source indication",
    101:"STATUS - Status",102:"INFORM - Information",
    103:"TOPSHP - Topmark/daymark shape",104:"TRAFCF - Traffic flow",
    105:"VALMAG - Local magnetic anomaly value",
    106:"VALMAG - Magnetic variation value",107:"VALSOU - Value of sounding",
    108:"VERACC - Vertical accuracy",109:"VERDAT - Vertical datum",
    110:"VERLEN - Vertical length",111:"WATLEV - Water level effect",
    112:"CATZOC - Category of zone of confidence",
    113:"DRVAL1 - Depth range value 1",114:"DRVAL2 - Depth range value 2",
    115:"EXPSOU - Exposition of sounding",
    116:"OBJNAM - Object name",
    117:"QUASOU - Quality of sounding measurement",
    118:"SCAMAX - Scale maximum",119:"SCAMIN - Scale minimum",
    120:"TECSOU - Technique of sounding measurement",
    121:"TXTDSO - Textual description of object",
    125:"QUASOU - Quality of sounding measurement",
    147:"SORDAT - Source date",148:"SORIND - Source indication",
    187:"WATLEV - Water level effect",
    300:"NINFOM - Information in national language",
    301:"NOBJNM - Object name in national language",
    302:"TXTDSC - Textual description in national language",
}

ATTVAL = {
    "CATWRK":{1:"Dangerous wreck",2:"Non-dangerous wreck",
              3:"Distributed remains of wreck",4:"Wreck showing mast/masts",
              5:"Wreck showing any portion of hull/superstructure"},
    "QUASOU":{1:"Depth known",2:"Depth unknown",3:"Doubtful sounding",
              4:"Unreliable sounding",5:"No bottom"},
    "WATLEV":{1:"Partly submerged at high water",2:"Always dry",
              3:"Always under water/submerged",4:"Covers and uncovers",
              5:"Awash",6:"Subject to inundation or flooding",7:"Floating"},
}

# DDR field structure tree (parent-child relationships per reference doc)
# tag -> (parent_tag, is_single_or_multi, field_name)
DDR_FIELD_INFO = {
    '0001': (None, 'single', 'ISO 8211 Record Identifier'),
    'DSID': (None, 'single', 'Data set identification field'),
    'DSSI': (None, 'single', 'Data set structure information field'),
    'DSPM': (None, 'single', 'Data set parameter field'),
    'FRID': (None, 'single', 'Feature record identifier field'),
    'FOID': (None, 'single', 'Feature object identifier field'),
    'ATTF': (None, 'multi',  'Feature record attribute field'),
    'NATF': (None, 'multi',  'Feature record national attribute field'),
    'FFPT': (None, 'multi',  'Feature record to feature object pointer field'),
    'FSPT': (None, 'multi',  'Feature record to spatial record pointer field'),
    'VRID': (None, 'single', 'Vector record identifier field'),
    'ATTV': (None, 'multi',  'Vector record attribute field'),
    'VRPT': (None, 'multi',  'Vector record pointer field'),
    'SG2D': (None, 'multi',  '2-D Coordinate field'),
    'SG3D': (None, 'multi',  '3-D Coordinate field'),
}
# Parent-child relationships from DDR field area
# value = parent tag; key = child tag
DDR_TREE_PARENTS = {
    'DSSI': 'DSID',
    'FOID': 'FRID', 'ATTF': 'FRID', 'NATF': 'FRID', 'FFPT': 'FRID', 'FSPT': 'FRID',
    'ATTV': 'VRID', 'VRPT': 'VRID', 'SG2D': 'VRID', 'SG3D': 'VRID',
}

# Subfield definitions per field tag
DDR_SUBFIELDS = {
    '0001': [('Record Identifier', 'b12')],
    'DSID': [('RCNM','b11'),('RCID','b14'),('EXPP','b11'),('INTU','b11'),
             ('DSNM','A'),('EDTN','A'),('UPDN','A'),('UADT','A(8)'),('ISDT','A(8)'),
             ('STED','R(4)'),('PRSP','b11'),('PSDN','A'),('PRED','A'),
             ('PROF','b11'),('AGEN','b12'),('COMT','A')],
    'DSSI': [('DSTR','b11'),('AALL','b11'),('NALL','b11'),('NOMR','b14'),
             ('NOCR','b14'),('NOGR','b14'),('NOLR','b14'),('NOIN','b14'),
             ('NOCN','b14'),('NOED','b14'),('NOFA','b14')],
    'DSPM': [('RCNM','b11'),('RCID','b14'),('HDAT','b11'),('VDAT','b11'),('SDAT','b11'),
             ('CSCL','b14'),('DUNI','b11'),('HUNI','b11'),('PUNI','b11'),
             ('COUN','b11'),('COMF','b14'),('SOMF','b14'),('COMT','A')],
    'FRID': [('RCNM','b11'),('RCID','b14'),('PRIM','b11'),('GRUP','b11'),
             ('OBJL','b12'),('RVER','b12'),('RUIN','b11')],
    'FOID': [('AGEN','b12'),('FIDN','b14'),('FIDS','b12')],
    'ATTF': [('ATTL','b12'),('ATVL','A')],
    'NATF': [('ATTL','b12'),('ATVL','A')],
    'FFPT': [('LNAM','B(64)'),('RIND','b11'),('COMT','A')],
    'FSPT': [('NAME','B(40)'),('ORNT','b11'),('USAG','b11'),('MASK','b11')],
    'VRID': [('RCNM','b11'),('RCID','b14'),('RVER','b12'),('RUIN','b11')],
    'ATTV': [('ATTL','b12'),('ATVL','A')],
    'VRPT': [('NAME','B(40)'),('ORNT','b11'),('USAG','b11'),('TOPI','b11'),('MASK','b11')],
    'SG2D': [('YCOO','b24'),('XCOO','b24')],
    'SG3D': [('YCOO','b24'),('XCOO','b24'),('VE3D','b24')],
}

# ============================================================
# ISO 8211 CORE
# ============================================================
def parse_leader(raw):
    """Parse 24-byte ISO 8211 leader."""
    rec_len = int(raw[0:5])
    ba_str = raw[12:17].decode('ascii', errors='replace').strip()
    base_addr = int(ba_str) if ba_str.isdigit() else 0
    def _em(b):
        c = chr(b); return int(c) if c.isdigit() else 0
    return {
        'record_length': rec_len,
        'interchange_level': chr(raw[5]),
        'leader_id': chr(raw[6]),
        'inline_code_ext': chr(raw[7]),
        'version': chr(raw[8]),
        'app_indicator': chr(raw[9]),
        'field_ctrl_len': raw[10:12].decode('ascii', errors='replace'),
        'base_addr': base_addr,
        'ext_char_set': raw[17:20].decode('ascii', errors='replace'),
        'entry_map': {'len_size':_em(raw[20]),'pos_size':_em(raw[21]),
                       'reserved':_em(raw[22]),'tag_size':_em(raw[23])}
    }

def parse_directory(dir_raw, em):
    """Parse directory entries."""
    ts, ls, ps = em['tag_size'], em['len_size'], em['pos_size']
    if ts<=0 or ls<=0 or ps<=0: return []
    esz = ts + ls + ps
    entries = []
    for i in range(0, len(dir_raw)-esz+1, esz):
        e = dir_raw[i:i+esz]
        try:
            tag = e[:ts].decode('ascii').strip()
            flen = int(e[ts:ts+ls])
            fpos = int(e[ts+ls:ts+ls+ps])
            if flen>0: entries.append({'tag':tag,'length':flen,'position':fpos})
        except: continue
    return entries

# ============================================================
# BINARY FIELD PARSERS
# ============================================================
def parse_dsid(data):
    """DSID: RCNM(1LE)+RCID(4LE)+EXPP(1)+INTU(1)+DSNM(A)+EDTN(A)+UPDN(A)+
       UADT(8A)+ISDT(8A)+STED(4 raw hex)+PRSP(1)+PSDN(A)+PRED(A)+PROF(1)+AGEN(2LE)+COMT(A)"""
    d = data.rstrip(b'\x1e')
    r = {}; p = 0
    try:
        r['RCNM']=d[p]; p+=1
        r['RCID']=int.from_bytes(d[p:p+4],'little'); p+=4
        r['EXPP']=d[p]; p+=1
        r['INTU']=d[p]; p+=1
        # DSNM: ASCII until 0x1F
        end=d.find(b'\x1f',p)
        r['DSNM']=d[p:end].decode('ascii',errors='replace').strip() if end>=0 else d[p:].decode('ascii',errors='replace').strip(); p=end+1 if end>=0 else len(d)
        end=d.find(b'\x1f',p)
        r['EDTN']=d[p:end].decode('ascii',errors='replace').strip() if end>=0 else d[p:].decode('ascii',errors='replace').strip(); p=end+1 if end>=0 else len(d)
        end=d.find(b'\x1f',p)
        r['UPDN']=d[p:end].decode('ascii',errors='replace').strip() if end>=0 else d[p:].decode('ascii',errors='replace').strip(); p=end+1 if end>=0 else len(d)
        # UADT/ISDT: fixed 8 ASCII
        r['UADT']=d[p:p+8].decode('ascii',errors='replace').strip('\x00 '); p+=8
        r['ISDT']=d[p:p+8].decode('ascii',errors='replace').strip('\x00 '); p+=8
        # STED: R(4) = 4 bytes raw hex
        r['STED']=d[p:p+4].hex(); p+=4
        r['PRSP']=d[p]; p+=1
        end=d.find(b'\x1f',p)
        r['PSDN']=d[p:end].decode('ascii',errors='replace').strip() if end>=0 else d[p:].decode('ascii',errors='replace').strip(); p=end+1 if end>=0 else len(d)
        end=d.find(b'\x1f',p)
        r['PRED']=d[p:end].decode('ascii',errors='replace').strip() if end>=0 else d[p:].decode('ascii',errors='replace').strip(); p=end+1 if end>=0 else len(d)
        r['PROF']=d[p]; p+=1
        r['AGEN']=int.from_bytes(d[p:p+2],'little'); p+=2
        # Remaining is COMT
        end=d.find(b'\x1f',p)
        r['COMT']=d[p:end].decode('ascii',errors='replace').strip() if end>=0 else d[p:].decode('ascii',errors='replace').strip() if p<len(d) else ''
    except Exception as e:
        r['_error']=str(e)
    return r

def parse_dssi(data):
    """DSSI: DSTR(1)+AALL(1)+NALL(1)+NOMR(4LE)+NOCR(4LE)+NOGR(4LE)+
       NOLR(4LE)+NOIN(4LE)+NOCN(4LE)+NOED(4LE)+NOFA(4LE)"""
    d = data.rstrip(b'\x1e')
    r = {}; p = 0
    names=['DSTR','AALL','NALL','NOMR','NOCR','NOGR','NOLR','NOIN','NOCN','NOED','NOFA']
    sizes=[1,1,1,4,4,4,4,4,4,4,4]
    for n,sz in zip(names,sizes):
        if p+sz>len(d): break
        r[n]=int.from_bytes(d[p:p+sz],'little'); p+=sz
    return r

def parse_dspm(data):
    """DSPM: RCNM(1)+RCID(4LE)+HDAT(1)+VDAT(1)+SDAT(1)+CSCL(4LE)+
       DUNI(1)+HUNI(1)+PUNI(1)+COUN(1)+COMF(4LE)+SOMF(4LE)+COMT(A)"""
    d = data.rstrip(b'\x1e')
    r = {}; p = 0
    try:
        r['RCNM']=d[p]; p+=1
        r['RCID']=int.from_bytes(d[p:p+4],'little'); p+=4
        r['HDAT']=d[p]; p+=1
        r['VDAT']=d[p]; p+=1
        r['SDAT']=d[p]; p+=1
        r['CSCL']=int.from_bytes(d[p:p+4],'little'); p+=4
        r['DUNI']=d[p]; p+=1
        r['HUNI']=d[p]; p+=1
        r['PUNI']=d[p]; p+=1
        r['COUN']=d[p]; p+=1
        r['COMF']=int.from_bytes(d[p:p+4],'little'); p+=4
        r['SOMF']=int.from_bytes(d[p:p+4],'little'); p+=4
        if p<len(d): r['COMT']=d[p:].rstrip(b'\x1e').decode('ascii',errors='replace').strip()
    except Exception as e: r['_error']=str(e)
    return r

def parse_frid(data):
    """FRID: RCNM(1)+RCID(4LE)+PRIM(1)+GRUP(1)+OBJL(2LE)+RVER(2LE)+RUIN(1)"""
    d=data.rstrip(b'\x1e')
    if len(d)>=12:
        return {'RCNM':d[0],'RCID':int.from_bytes(d[1:5],'little'),'PRIM':d[5],
                'GRUP':d[6],'OBJL':int.from_bytes(d[7:9],'little'),
                'RVER':int.from_bytes(d[9:11],'little'),'RUIN':d[11]}
    return {}

def parse_foid(data):
    """FOID: AGEN(2LE)+FIDN(4LE)+FIDS(2LE) per instance."""
    d=data.rstrip(b'\x1e')
    inst=[]
    for i in range(0,len(d)-7,8):
        b=d[i:i+8]
        if len(b)==8: inst.append({'AGEN':int.from_bytes(b[0:2],'little'),
            'FIDN':int.from_bytes(b[2:6],'little'),'FIDS':int.from_bytes(b[6:8],'little')})
    return inst

def parse_attf(data):
    """ATTF: ATTL(2LE)+ATVL(ASCII until 0x1F or 0x1E)."""
    d=data.rstrip(b'\x1e'); attrs=[]; p=0
    while p+2<=len(d):
        attl=int.from_bytes(d[p:p+2],'little'); p+=2
        if p>=len(d): break
        end=d.find(b'\x1f',p)
        if end<0: atvl_bytes=d[p:]; p=len(d)
        else: atvl_bytes=d[p:end]; p=end+1
        try: atvl=atvl_bytes.decode('ascii',errors='replace').strip()
        except: atvl=atvl_bytes.hex(' ')
        attrs.append({'ATTL':attl,'ATTL_hex':f'0x{attl:04X}','ATVL':atvl})
    return attrs

def parse_natf(data):
    """NATF: ATTL(2LE)+ATVL(UTF-16LE until 0x1F).
    After 0x1F separator, a 0x00 alignment byte may follow.
    ATVL bytes may contain embedded 0x00 bytes from ASCII digits in UTF-16LE.
    We decode the full ATVL as UTF-16LE, stripping trailing null bytes.
    """
    d=data.rstrip(b'\x1e'); attrs=[]; p=0
    while p+2<=len(d):
        attl=int.from_bytes(d[p:p+2],'little'); p+=2
        if p>=len(d): break
        end=d.find(b'\x1f',p)
        if end<0: atvl_bytes=d[p:]; p=len(d)
        else: atvl_bytes=d[p:end]; p=end+1
        # Skip 0x00 alignment byte after separator
        if p<len(d) and d[p]==0x00: p+=1
        # Trim trailing nulls then decode as UTF-16LE
        trimmed=atvl_bytes.rstrip(b'\x00')
        if len(trimmed)%2!=0: trimmed+=(b'\x00')  # ensure even length
        try: atvl=trimmed.decode('utf-16-le',errors='replace')
        except: atvl=trimmed.hex(' ')
        attrs.append({'ATTL':attl,'ATTL_hex':f'0x{attl:04X}','ATVL':atvl})
    return attrs

def parse_fspt(data):
    """FSPT: NAME(5B=raw hex)+ORNT(1)+USAG(1)+MASK(1)."""
    d=data.rstrip(b'\x1e'); entries=[]
    for i in range(0,len(d)-7,8):
        b=d[i:i+8]
        if len(b)==8:
            entries.append({'NAME_hex':b[0:5].hex(),'RCNM':b[0],
                'RCID_BE':int.from_bytes(b[1:5],'big'),'ORNT':b[5],'USAG':b[6],'MASK':b[7]})
    return entries

def parse_vrid(data):
    """VRID: RCNM(1)+RCID(4LE)+RVER(2LE)+RUIN(1)"""
    d=data.rstrip(b'\x1e')
    if len(d)>=8:
        return {'RCNM':d[0],'RCID':int.from_bytes(d[1:5],'little'),
                'RVER':int.from_bytes(d[5:7],'little'),'RUIN':d[7]}
    return {}

def parse_sg2d(data,comf=10000000):
    """SG2D: YCOO+XCOO. Auto-detect 3-byte vs 4-byte per value."""
    d=data.rstrip(b'\x1e'); coords=[]
    stride=8 if len(d)>=8 and len(d)%8 in(0,1) else 6
    vb=stride//2
    for i in range(0,len(d)-stride+1,stride):
        coords.append((int.from_bytes(d[i:i+vb],'little',signed=True)/comf,
                       int.from_bytes(d[i+vb:i+2*vb],'little',signed=True)/comf))
    return coords

# ============================================================
# REPORT HELPERS
# ============================================================
def xml_safe(text):
    if not isinstance(text,str): text=str(text)
    result=[]
    for ch in text:
        cp=ord(ch)
        if cp in(0x09,0x0A,0x0D) or (0x20<=cp<=0xD7FF) or (0xE000<=cp<=0xFFFD): result.append(ch)
        elif cp<0x20: result.append(' ')
    return ''.join(result)

def safe_print(s):
    try: print(s)
    except UnicodeEncodeError:
        try: print(s.encode('utf-8',errors='replace').decode('utf-8'))
        except: print(s.encode('ascii',errors='replace').decode('ascii'))

def get_attr_label(attl):
    """Get 'Attribute' and 'Acronym' columns from ATTR_NAMES."""
    name=ATTR_NAMES.get(attl,'')
    if ' - ' in name:
        acro=name.split(' - ')[0].strip()
        label=name.split(' - ')[1].strip()
    else:
        acro=''; label=name
    return label,acro

# ============================================================
# MAIN
# ============================================================
def main():
    global ENC_FILE
    ENC_FILE = get_enc_file_path()

    print("="*70)
    print("  S-57 ENC File Parser - Group 3 (第3小组)")
    print(f"  Target FOID: CN {TARGET_FIDN:010d} {TARGET_FIDS}")
    print("="*70)

    with open(ENC_FILE,'rb') as f: data=f.read()
    print(f"\nFile: {os.path.basename(ENC_FILE)} ({len(data)} bytes)")

    # Iterate through records
    offset=0; ddr_len=0; ddr_leader=None; ddr_dir=None; ddr_field_area=None; ddr_fields={}
    first_dr_fields={}; co2mf=10000000
    target_info=None; spatial_info=None
    records_parsed=0

    while offset<len(data) and records_parsed<15000:
        if offset+24>len(data): break
        leader=parse_leader(data[offset:offset+24])
        rlen=leader['record_length']; ba=leader['base_addr']
        if rlen<=0 or offset+rlen>len(data) or rlen>100000: break

        dir_raw=data[offset+24:offset+ba]
        field_area=data[offset+ba:offset+rlen]
        directory=parse_directory(dir_raw,leader['entry_map'])
        fields={}
        for e in directory:
            p,l=e['position'],e['length']
            if p+l<=len(field_area): fields[e['tag']]=field_area[p:p+l]

        records_parsed+=1

        # --- Record #1: DDR ---
        if records_parsed==1:
            ddr_leader=leader; ddr_dir=directory; ddr_field_area=field_area
            ddr_len=leader['record_length']
            for e in directory:
                p,l=e['position'],e['length']
                if p+l<=len(field_area): ddr_fields[e['tag']]=field_area[p:p+l]

            print(f"\n{'='*70}")
            print("RECORD #1: DDR (Data Descriptive Record)")
            print(f"{'='*70}")
            em=leader['entry_map']
            print(f"\n[Leader / DDR Header Area]")
            for k,v in [('Record Length',str(leader['record_length'])),
                        ('Interchange Level',leader['interchange_level']),
                        ('Leader ID',leader['leader_id']+' (DDR)'),
                        ('Inline Code Extension',leader['inline_code_ext']),
                        ('Version',leader['version']),
                        ('Application Indicator',repr(leader['app_indicator'])),
                        ('Field Control Length',leader['field_ctrl_len']),
                        ('Base Address',str(leader['base_addr'])),
                        ('Extended Char Set',repr(leader['ext_char_set'])),
                        ('Field Length Size',str(em['len_size'])),
                        ('Field Position Size',str(em['pos_size'])),
                        ('Reserved',str(em['reserved'])),
                        ('Field Tag Size',str(em['tag_size']))]:
                safe_print(f"  {k}: {v}")

            print(f"\n[DDR Directory / Table of Contents]")
            for i,e in enumerate(directory):
                safe_print(f"  #{i+1}: Tag={e['tag']}, Length={e['length']}, Position={e['position']}")

            # Build tree ordering (used by both console & DOCX)
            dir_tags=[e['tag'] for e in directory]
            ordered_tree=[]
            for tag in dir_tags:
                if tag in DDR_TREE_PARENTS: continue  # skip children at top level
                ordered_tree.append((tag,0))
                for ct in dir_tags:
                    if DDR_TREE_PARENTS.get(ct)==tag:
                        ordered_tree.append((ct,1))
                        for gc in dir_tags:
                            if DDR_TREE_PARENTS.get(gc)==ct:
                                ordered_tree.append((gc,2))

            print(f"\n[DDR Directory Tree Structure]")
            printed=set()
            for tag,level in ordered_tree:
                indent='    '*level
                info=DDR_FIELD_INFO.get(tag,(None,'?','?'))
                safe_print(f"{indent}[{tag}] {info[2]}")
                printed.add(tag)

            print(f"\n[DDR Field Descriptions]")
            for e in directory:
                tag=e['tag']
                info=DDR_FIELD_INFO.get(tag,(None,'single',''))
                fdata=ddr_fields.get(tag,b'')
                try: desc=fdata.rstrip(b'\x1e').decode('ascii',errors='replace')
                except: desc=fdata[:60].hex(' ')
                safe_print(f"  [{tag}] {info[1]}: {desc[:250]}")
                subs=DDR_SUBFIELDS.get(tag,[])
                for sf_name,sf_fmt in subs:
                    safe_print(f"      {sf_name} ({sf_fmt})")

        # --- Record #2: First DR (DSID+DSSI) ---
        elif records_parsed==2:
            print(f"\n{'='*70}")
            print("RECORD #2: First DR (Data Record)")
            print(f"{'='*70}")
            safe_print(f"\n[Leader] len={leader['record_length']}, base={leader['base_addr']}")
            safe_print(f"[Directory]")
            for e in directory:
                safe_print(f"  {e['tag']}: len={e['length']}, pos={e['position']}")

            if 'DSID' in fields:
                dsid=parse_dsid(fields['DSID'])
                print(f"\n[DSID]")
                for k,v in dsid.items():
                    safe_print(f"  {k}: {v}")
                first_dr_fields['DSID']=dsid
            if 'DSSI' in fields:
                dssi=parse_dssi(fields['DSSI'])
                print(f"\n[DSSI]")
                for k,v in dssi.items():
                    safe_print(f"  {k}: {v}")
                first_dr_fields['DSSI']=dssi

        # --- Record #3+: Look for DSPM (COMF) and Target FOID ---
        else:
            # DSPM: get COMF
            if 'DSPM' in fields:
                dspm=parse_dspm(fields['DSPM'])
                co2mf=dspm.get('COMF',10000000)
                safe_print(f"\n[DSPM] COMF={co2mf}")
                first_dr_fields['DSPM']=dspm

            # FOID: check for target
            if 'FOID' in fields and target_info is None:
                foid_list=parse_foid(fields['FOID'])
                for foi in foid_list:
                    if foi['AGEN']==TARGET_AGENCY and foi['FIDN']==TARGET_FIDN and foi['FIDS']==TARGET_FIDS:
                        target_info={
                            'offset':offset,'leader':leader,'directory':directory,'fields':fields,
                            'frid':parse_frid(fields.get('FRID',b'')),
                            'foid_list':foid_list,
                            'attrs':parse_attf(fields.get('ATTF',b'')),
                            'natf_attrs':parse_natf(fields.get('NATF',b'')),
                            'fspt_list':parse_fspt(fields.get('FSPT',b'')),
                        }
                        print(f"\n{'='*70}")
                        print(f"*** TARGET FEATURE FOUND at offset {offset} ***")
                        print(f"{'='*70}")
                        frid=target_info['frid']
                        print(f"\n[FRID - Feature Record Identifier]")
                        for k,v in frid.items():
                            desc=''
                            if k=='RCNM': desc=RCNM_NAMES.get(v,'')
                            elif k=='OBJL': desc=OBJL_NAMES.get(v,'')
                            safe_print(f"  {k}: {v}{' ('+desc+')' if desc else ''}")
                        print(f"\n[FOID - Feature Object Identifier]")
                        for foi in foid_list:
                            safe_print(f"  AGEN={foi['AGEN']} (CN), FIDN={foi['FIDN']:010d}, FIDS={foi['FIDS']}")
                        print(f"\n[FOID Encoding (Little-Endian)]")
                        safe_print(f"  CN=70 -> 0x0046 -> bytes: 46 00")
                        safe_print(f"  0013753099=13753099=0x00D1DB0B -> bytes: 0B DB D1 00")
                        safe_print(f"  22501=0x57E5 -> bytes: E5 57")
                        safe_print(f"  Complete: 46 00 0B DB D1 00 E5 57")

                        print(f"\n[ATTF - Feature Record Attributes]")
                        for a in target_info['attrs']:
                            attl=a['ATTL']; atvl=a['ATVL']
                            label,acro=get_attr_label(attl)
                            # Category value
                            val_str=str(atvl)
                            for cat_key,cat_dict in ATTVAL.items():
                                for code,name in ATTR_NAMES.items():
                                    if code==attl and cat_key in name.upper():
                                        if atvl.isdigit(): cv=cat_dict.get(int(atvl),''); val_str=f'{atvl} ({cv})' if cv else str(atvl)
                                        break
                            safe_print(f"  ATTL={attl} ({label}): ATVL={val_str}")

                        print(f"\n[NATF - National Language Attributes]")
                        for a in target_info['natf_attrs']:
                            attl=a['ATTL']; atvl=a['ATVL']
                            label,acro=get_attr_label(attl)
                            safe_print(f"  ATTL={attl} ({label}): {atvl}")

                        print(f"\n[FSPT - Feature to Spatial Record Pointer]")
                        for fsp in target_info['fspt_list']:
                            safe_print(f"  NAME (hex): {fsp['NAME_hex']}")
                            safe_print(f"    RCNM={fsp['RCNM']} ({RCNM_NAMES.get(fsp['RCNM'],'')})")
                            safe_print(f"    RCID (BE)=0x{fsp['RCID_BE']:08X}")
                            safe_print(f"    ORNT={fsp['ORNT']}, USAG={fsp['USAG']}, MASK={fsp['MASK']}")

                        # Convert FSPT RCID_BE -> RCID_LE for VRID lookup
                        for fsp in target_info['fspt_list']:
                            rcid_bytes=fsp['RCID_BE'].to_bytes(4,'big')
                            rcid_le=int.from_bytes(rcid_bytes,'little')
                            vname=f'{fsp["RCNM"]:03d}{rcid_le:08d}'
                            print(f"\n{'='*70}")
                            print(f"SEARCHING FOR VRID: {vname}")
                            print(f"{'='*70}")

                            # Scan for matching VRID (start from after DDR)
                            off2=ddr_len
                            found=False; scans=0
                            while off2<len(data) and scans<15000 and not found:
                                if off2+24>len(data): break
                                if not data[off2:off2+5].isdigit(): break
                                rl2=int(data[off2:off2+5])
                                if rl2<=0 or rl2>100000: break
                                if chr(data[off2+6]) not in('L','D'): off2+=rl2; scans+=1; continue
                                bs2=data[off2+12:off2+17].decode(errors='replace').strip()
                                if not bs2.isdigit(): off2+=rl2; scans+=1; continue
                                ba2=int(bs2)
                                ldr2=parse_leader(data[off2:off2+24])
                                dir2=parse_directory(data[off2+24:off2+ba2],ldr2['entry_map'])
                                fa2=data[off2+ba2:off2+rl2]
                                for e2 in dir2:
                                    if e2['tag']=='VRID' and e2['position']+e2['length']<=len(fa2):
                                        vrid=parse_vrid(fa2[e2['position']:e2['position']+e2['length']])
                                        vn2=f'{vrid["RCNM"]:03d}{vrid["RCID"]:08d}'
                                        if vn2==vname:
                                            spatial_info={'name':vn2,'vrid':vrid,'offset':off2}
                                            print(f"\n  FOUND at offset {off2}!")
                                            print(f"\n  [VRID - Vector Record Identifier]")
                                            for k,v in vrid.items():
                                                if k=='RCNM': safe_print(f"    {k}: {v} ({RCNM_NAMES.get(v,'')})")
                                                else: safe_print(f"    {k}: {v}")
                                            for e2 in dir2:
                                                if e2['tag'] in('SG2D','SG3D') and e2['position']+e2['length']<=len(fa2):
                                                    if e2['tag']=='SG2D':
                                                        coords=parse_sg2d(fa2[e2['position']:e2['position']+e2['length']],co2mf)
                                                        spatial_info['coord_type']='2D'
                                                    else:
                                                        coords=parse_sg2d(fa2[e2['position']:e2['position']+e2['length']],co2mf)
                                                        spatial_info['coord_type']='3D'
                                                    spatial_info['coords']=coords
                                                    print(f"\n  [{e2['tag']} - Coordinates] ({len(coords)} points, COMF={co2mf})")
                                                    for ci,c in enumerate(coords[:20]):
                                                        safe_print(f"    [{ci+1}] YCOO(Lat)={c[0]:.7f}, XCOO(Lon)={c[1]:.7f}")
                                                    if len(coords)>20: safe_print(f"    ... ({len(coords)-20} more)")
                                            found=True
                                off2+=rl2; scans+=1
                            if not found:
                                safe_print(f"  NOT FOUND (may be external reference)")

        offset+=rlen

    # ================================================================
    # GENERATE DOCX REPORT
    # ================================================================
    print(f"\n{'='*70}")
    print("GENERATING WORD REPORT (DOCX)")
    print(f"{'='*70}")

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc=Document()
        doc.styles['Normal'].font.size=Pt(10)

        h=doc.add_heading('S-57 ENC File Parsing Report',0)
        h.alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Group 3 / Di 3 Xiao Zu')
        doc.add_paragraph(f'File: {os.path.basename(ENC_FILE)}')
        doc.add_paragraph(f'Target FOID: CN {TARGET_FIDN:010d} {TARGET_FIDS}')
        doc.add_paragraph('Date: 2026-06-28')

        # === 1. DDR Leader ===
        doc.add_heading('1. DDR Header Area',1)
        ld=[('Record Length',str(ddr_leader['record_length'])),
            ('Interchange Level',ddr_leader['interchange_level']),
            ('Leader ID',ddr_leader['leader_id']),
            ('Inline Code Extension',ddr_leader['inline_code_ext']),
            ('Version',ddr_leader['version']),
            ('Application Indicator',ddr_leader['app_indicator']),
            ('Field Control Length',ddr_leader['field_ctrl_len']),
            ('Base Address',str(ddr_leader['base_addr'])),
            ('Extended Char Set',ddr_leader['ext_char_set']),
            ('Field Length Size',str(ddr_leader['entry_map']['len_size'])),
            ('Field Position Size',str(ddr_leader['entry_map']['pos_size'])),
            ('Reserved',str(ddr_leader['entry_map']['reserved'])),
            ('Field Tag Size',str(ddr_leader['entry_map']['tag_size']))]
        t=doc.add_table(rows=len(ld)+1,cols=2,style='Table Grid')
        t.cell(0,0).text='Field'; t.cell(0,1).text='Value'
        for i,(k,v) in enumerate(ld): t.cell(i+1,0).text=k; t.cell(i+1,1).text=str(v)

        # === 2. DDR Directory ===
        doc.add_heading('2. DDR Directory',1)
        dt=doc.add_table(rows=len(ddr_dir)+1,cols=3,style='Table Grid')
        dt.cell(0,0).text='Field Tag'; dt.cell(0,1).text='Length'; dt.cell(0,2).text='Offset'
        for i,e in enumerate(ddr_dir): dt.cell(i+1,0).text=e['tag']; dt.cell(i+1,1).text=str(e['length']); dt.cell(i+1,2).text=str(e['position'])

        # --- Helper for counting table rows ---
        def _sf_rows(tag):
            s=DDR_SUBFIELDS.get(tag,[])
            return max(0,len(s)-1)

        # === 3. DDR Tree (exact reference layout) ===
        doc.add_heading('3. DDR Directory Tree Structure',1)
        # Count rows: 1 per tag + extra for subfields beyond the first
        tree_row_count=sum(1+_sf_rows(tag) for tag,_ in ordered_tree)
        tt=doc.add_table(rows=tree_row_count+1,cols=5,style='Table Grid')
        tt.cell(0,0).text='Field Tag'; tt.cell(0,1).text='Repeat?'; tt.cell(0,2).text='Field Name'
        tt.cell(0,3).text='Sub-field Name'; tt.cell(0,4).text='Sub-field Type'
        ri=1
        for tag,level in ordered_tree:
            info=DDR_FIELD_INFO.get(tag,(None,'?','?'))
            indent='    '*level
            tt.cell(ri,0).text=indent+tag
            tt.cell(ri,1).text=info[1]
            tt.cell(ri,2).text=info[2]
            subs=DDR_SUBFIELDS.get(tag,[])
            if subs: tt.cell(ri,3).text=subs[0][0]; tt.cell(ri,4).text=subs[0][1]
            ri+=1
            for sf in subs[1:]:
                tt.cell(ri,3).text=sf[0]; tt.cell(ri,4).text=sf[1]; ri+=1

        # === 4. DDR Field Descriptions ===
        doc.add_heading('4. DDR Field Descriptions',1)
        fd_row_count=sum(1+_sf_rows(tag) for tag,_ in ordered_tree)
        ft=doc.add_table(rows=fd_row_count+1,cols=5,style='Table Grid')
        ft.cell(0,0).text='Field Tag'; ft.cell(0,1).text='Repeat?'; ft.cell(0,2).text='Field Name'
        ft.cell(0,3).text='Sub-field Name'; ft.cell(0,4).text='Sub-field Type'
        ri=1
        for tag,_level in ordered_tree:
            info=DDR_FIELD_INFO.get(tag,(None,'?','?'))
            ft.cell(ri,0).text=tag; ft.cell(ri,1).text=info[1]; ft.cell(ri,2).text=info[2]
            subs=DDR_SUBFIELDS.get(tag,[])
            if subs: ft.cell(ri,3).text=subs[0][0]; ft.cell(ri,4).text=subs[0][1]
            ri+=1
            for sf in subs[1:]:
                ft.cell(ri,3).text=sf[0]; ft.cell(ri,4).text=sf[1]; ri+=1

        # === 5. First DR ===
        doc.add_heading('5. First DR (Data Record)',1)
        if 'DSID' in first_dr_fields:
            dsid=first_dr_fields['DSID']
            ds=doc.add_table(rows=len(dsid)+1,cols=4,style='Table Grid')
            ds.cell(0,0).text='Field Tag'; ds.cell(0,1).text='Instance#'
            ds.cell(0,2).text='Sub-field Name'; ds.cell(0,3).text='Value'
            for i,(k,v) in enumerate(dsid.items()):
                ds.cell(i+1,0).text='DSID'; ds.cell(i+1,1).text='0'
                ds.cell(i+1,2).text=k; ds.cell(i+1,3).text=xml_safe(str(v))
        if 'DSSI' in first_dr_fields:
            dssi=first_dr_fields['DSSI']
            di=doc.add_table(rows=len(dssi)+1,cols=4,style='Table Grid')
            di.cell(0,0).text='Field Tag'; di.cell(0,1).text='Instance#'
            di.cell(0,2).text='Sub-field Name'; di.cell(0,3).text='Value'
            for i,(k,v) in enumerate(dssi.items()):
                di.cell(i+1,0).text='DSSI'; di.cell(i+1,1).text='0'
                di.cell(i+1,2).text=k; di.cell(i+1,3).text=str(v)

        # === 6. DSPM / COMF ===
        doc.add_heading('6. DSPM (Data Set Parameters)',1)
        ct=doc.add_table(rows=2,cols=2,style='Table Grid')
        ct.cell(0,0).text='Field'; ct.cell(0,1).text='Value'
        ct.cell(1,0).text='COMF'; ct.cell(1,1).text=str(co2mf)

        # === 7. Target Feature Record ===
        if target_info:
            doc.add_heading('7. Target Feature Record (Group 3)',1)
            doc.add_paragraph(f'Target FOID: CN {TARGET_FIDN:010d} {TARGET_FIDS}')
            doc.add_paragraph(f'Record offset: {target_info["offset"]} (0x{target_info["offset"]:06X})')

            # FOID encoding
            doc.add_heading('7.1 FOID Encoding (Little-Endian)',2)
            doc.add_paragraph('AGEN: CN=70=0x0046, LE bytes: 46 00')
            doc.add_paragraph('FIDN: 13753099=0x00D1DB0B, LE bytes: 0B DB D1 00')
            doc.add_paragraph('FIDS: 22501=0x57E5, LE bytes: E5 57')
            doc.add_paragraph('Complete: 46 00 0B DB D1 00 E5 57')

            # FRID
            frid=target_info['frid']
            doc.add_heading('7.2 FRID - Feature Record Identifier',2)
            ft2=doc.add_table(rows=len(frid)+1,cols=5,style='Table Grid')
            ft2.cell(0,0).text='Field Tag'; ft2.cell(0,1).text='Sub-field Name'
            ft2.cell(0,2).text='Value'; ft2.cell(0,3).text='Attribute'; ft2.cell(0,4).text='Acronym'
            for ri,(k,v) in enumerate(frid.items()):
                ft2.cell(ri+1,0).text='FRID'; ft2.cell(ri+1,1).text=k
                ft2.cell(ri+1,2).text=str(v)
                if k=='RCNM': ft2.cell(ri+1,3).text=RCNM_NAMES.get(v,'')
                elif k=='OBJL': ft2.cell(ri+1,3).text=OBJL_NAMES.get(v,'')
            # FOID
            for foi in target_info['foid_list']:
                foid_table=doc.add_table(rows=5,cols=5,style='Table Grid')
                foid_table.cell(0,0).text='Field Tag'; foid_table.cell(0,1).text='Sub-field Name'
                foid_table.cell(0,2).text='Value'; foid_table.cell(0,3).text='Attribute'; foid_table.cell(0,4).text='Acronym'
                for ri,(ft,sf,v,attr,acro) in enumerate([
                    ('FOID','AGEN',str(foi['AGEN']),'Production Agency','CN'),
                    ('FOID','FIDN',str(foi['FIDN']),'Feature Identification Number',''),
                    ('FOID','FIDS',str(foi['FIDS']),'Feature Identification Subdivision',''),
                    ('FOID','Legacy String',f'CN {foi["FIDN"]:010d} {foi["FIDS"]}','','')]):
                    foid_table.cell(ri+1,0).text=ft; foid_table.cell(ri+1,1).text=sf
                    foid_table.cell(ri+1,2).text=v; foid_table.cell(ri+1,3).text=attr; foid_table.cell(ri+1,4).text=acro

            # ATTF
            doc.add_heading('7.3 ATTF - Feature Record Attributes',2)
            attrs=target_info['attrs']
            if attrs:
                at=doc.add_table(rows=len(attrs)+1,cols=5,style='Table Grid')
                at.cell(0,0).text='Field Tag'; at.cell(0,1).text='Sub-field Name'
                at.cell(0,2).text='Value'; at.cell(0,3).text='Attribute'; at.cell(0,4).text='Acronym'
                for ri,a in enumerate(attrs):
                    label,acro=get_attr_label(a['ATTL'])
                    val_str=a['ATVL']
                    for cat_key,cat_dict in ATTVAL.items():
                        for code,name in ATTR_NAMES.items():
                            if code==a['ATTL'] and cat_key in name.upper() and a['ATVL'].isdigit():
                                cv=cat_dict.get(int(a['ATVL']),''); val_str=f'{a["ATVL"]} ({cv})' if cv else a['ATVL']
                                break
                    at.cell(ri+1,0).text='ATTF'; at.cell(ri+1,1).text=str(a['ATTL'])
                    at.cell(ri+1,2).text=xml_safe(val_str); at.cell(ri+1,3).text=label; at.cell(ri+1,4).text=acro

            # NATF
            doc.add_heading('7.4 NATF - National Language Attributes',2)
            natf_attrs=target_info['natf_attrs']
            if natf_attrs:
                nt=doc.add_table(rows=len(natf_attrs)+1,cols=5,style='Table Grid')
                nt.cell(0,0).text='Field Tag'; nt.cell(0,1).text='Sub-field Name'
                nt.cell(0,2).text='Value'; nt.cell(0,3).text='Attribute'; nt.cell(0,4).text='Acronym'
                for ri,a in enumerate(natf_attrs):
                    label,acro=get_attr_label(a['ATTL'])
                    nt.cell(ri+1,0).text='NATF'; nt.cell(ri+1,1).text=str(a['ATTL'])
                    nt.cell(ri+1,2).text=xml_safe(a['ATVL']); nt.cell(ri+1,3).text=label; nt.cell(ri+1,4).text=acro

            # FSPT
            doc.add_heading('7.5 FSPT - Feature to Spatial Record Pointer',2)
            fspt_list=target_info['fspt_list']
            if fspt_list:
                fst=doc.add_table(rows=len(fspt_list)*4+1,cols=5,style='Table Grid')
                fst.cell(0,0).text='Field Tag'; fst.cell(0,1).text='Sub-field Name'
                fst.cell(0,2).text='Value'; fst.cell(0,3).text='Attribute'; fst.cell(0,4).text='Acronym'
                ri=1
                for fsp in fspt_list:
                    fst.cell(ri,0).text='FSPT'; fst.cell(ri,1).text='NAME'
                    fst.cell(ri,2).text=fsp['NAME_hex']; ri+=1
                    fst.cell(ri,0).text='FSPT'; fst.cell(ri,1).text='ORNT'
                    fst.cell(ri,2).text=str(fsp['ORNT']); ri+=1
                    fst.cell(ri,0).text='FSPT'; fst.cell(ri,1).text='USAG'
                    fst.cell(ri,2).text=str(fsp['USAG']); ri+=1
                    fst.cell(ri,0).text='FSPT'; fst.cell(ri,1).text='MASK'
                    fst.cell(ri,2).text=str(fsp['MASK']); ri+=1

            # === 8. Spatial Record ===
            doc.add_heading('8. Associated Spatial Vector Record',1)
            if spatial_info:
                vrid=spatial_info['vrid']
                doc.add_paragraph(f'VRID NAME: {spatial_info["name"]}')
                for k,v in vrid.items():
                    desc=RCNM_NAMES.get(v,'') if k=='RCNM' else ''
                    doc.add_paragraph(f'{k}: {v}{" ("+desc+")" if desc else ""}',style='List Bullet')
                coords=spatial_info.get('coords',[])
                if coords:
                    doc.add_paragraph(f'Coordinates ({len(coords)} points, COMF={co2mf}):')
                    ct2=doc.add_table(rows=min(len(coords)+1,21),cols=3,style='Table Grid')
                    ct2.cell(0,0).text='No.'; ct2.cell(0,1).text='YCOO (Latitude)'; ct2.cell(0,2).text='XCOO (Longitude)'
                    for ci,c in enumerate(coords[:20]):
                        ct2.cell(ci+1,0).text=str(ci+1); ct2.cell(ci+1,1).text=f'{c[0]:.7f}'; ct2.cell(ci+1,2).text=f'{c[1]:.7f}'
                    if len(coords)>20: doc.add_paragraph(f'... ({len(coords)-20} more)')
            else:
                doc.add_paragraph('No associated spatial vector record found in this file.')

        # === Save ===
        output_path=os.path.join(OUTPUT_DIR,'s57_parse_report.docx')
        doc.save(output_path)
        print(f"\nReport saved: {output_path} ({os.path.getsize(output_path)} bytes)")

    except ImportError:
        print("\npython-docx not installed. Install: pip install python-docx")
        # Write text report
        txt_path=os.path.join(OUTPUT_DIR,'s57_parse_report.txt')
        with open(txt_path,'w',encoding='utf-8') as f:
            f.write("S-57 ENC PARSING REPORT - Group 3\n")
            f.write("="*60+"\n\n")
            f.write(f"Target FOID: CN {TARGET_FIDN:010d} {TARGET_FIDS}\n\n")
            if target_info:
                f.write("FRID:\n")
                for k,v in target_info['frid'].items(): f.write(f"  {k}: {v}\n")
                f.write("\nATTF:\n")
                for a in target_info['attrs']: f.write(f"  ATTL={a['ATTL']}: {a['ATVL']}\n")
                f.write("\nNATF:\n")
                for a in target_info['natf_attrs']: f.write(f"  ATTL={a['ATTL']}: {a['ATVL']}\n")
                f.write("\nFSPT:\n")
                for fsp in target_info['fspt_list']: f.write(f"  NAME={fsp['NAME_hex']}, ORNT={fsp['ORNT']}\n")
            if spatial_info:
                f.write(f"\nSpatial: {spatial_info['name']}\n")
                for ci,c in enumerate(spatial_info.get('coords',[])[:10]):
                    f.write(f"  [{ci+1}] YCOO={c[0]:.7f}, XCOO={c[1]:.7f}\n")
        print(f"Text report: {txt_path}")

    print(f"\n{'='*70}")
    print("PARSING COMPLETE!")
    print(f"{'='*70}")

if __name__=='__main__':
    main()
